import urllib.request
import os

DB_URL = "https://github.com/ik1903846-web/kap-tarayici/raw/main/kap_data.db"
DB_PATH = os.environ.get("DB_PATH", "kap_data.db")

if not os.path.exists(DB_PATH):
    print("Veritabanı indiriliyor...")
    urllib.request.urlretrieve(DB_URL, DB_PATH)
    print("Veritabanı indirildi!")
from flask import Flask, render_template, jsonify, send_file, abort
import sqlite3
import os
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DB_PATH = os.environ.get("DB_PATH", "kap_data.db")
app = Flask(__name__)

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def db_query(sql, params=()):
    conn = get_db()
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/companies")
def api_companies():
    companies = db_query("""
        SELECT c.id, c.code, c.title, c.updated_at,
               COUNT(DISTINCT f.year) as year_count,
               COUNT(f.id) as data_count
        FROM companies c
        LEFT JOIN financials f ON c.id = f.company_id
        GROUP BY c.id ORDER BY c.code
    """)
    return jsonify(companies)

@app.route("/api/financial/<company_id>")
def api_financial(company_id):
    company = db_query("SELECT * FROM companies WHERE id=?", (company_id,))
    if not company:
        abort(404)
    rows = db_query("""
        SELECT year, quarter, table_type, label, value, currency
        FROM financials WHERE company_id=?
        ORDER BY year DESC, table_type, label
    """, (company_id,))
    tables = {}
    years_set = set()
    for row in rows:
        tt = row["table_type"]
        label = row["label"].strip("/").replace("/", " › ")
        year = row["year"]
        years_set.add(year)
        tables.setdefault(tt, {}).setdefault(label, {})[year] = row["value"]
    return jsonify({
        "company": company[0],
        "years": sorted(years_set, reverse=True),
        "tables": tables
    })

@app.route("/download/<company_id>/excel")
def download_excel(company_id):
    company = db_query("SELECT * FROM companies WHERE id=?", (company_id,))
    if not company:
        abort(404)
    c = company[0]
    rows = db_query("""
        SELECT year, table_type, label, value FROM financials
        WHERE company_id=? ORDER BY table_type, label, year DESC
    """, (company_id,))
    if not rows:
        abort(404)
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    tables = {}
    years_set = set()
    for row in rows:
        years_set.add(row["year"])
        tables.setdefault(row["table_type"], {}).setdefault(
            row["label"].strip("/"), {})[row["year"]] = row["value"]
    years = sorted(years_set, reverse=True)
    for sheet_name, data in tables.items():
        ws = wb.create_sheet(title=sheet_name[:31])
        ws.merge_cells(f"A1:{chr(65+len(years))}1")
        t = ws.cell(1, 1, f"{c['code']} - {c['title']}")
        t.font = Font(bold=True, color="FFFFFF", size=13)
        t.fill = PatternFill("solid", fgColor="1F4E79")
        t.alignment = Alignment(horizontal="center")
        ws.cell(2, 1, "Kalem").font = Font(bold=True, color="FFFFFF")
        ws.cell(2, 1).fill = PatternFill("solid", fgColor="2E75B6")
        for ci, year in enumerate(years, 2):
            cell = ws.cell(2, ci, str(year))
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="2E75B6")
            cell.alignment = Alignment(horizontal="right")
        for ri, (label, year_vals) in enumerate(sorted(data.items()), 3):
            ws.cell(ri, 1, label)
            for ci, year in enumerate(years, 2):
                val = year_vals.get(year)
                if val is not None:
                    c2 = ws.cell(ri, ci, val)
                    c2.number_format = '#,##0.00'
                    c2.alignment = Alignment(horizontal="right")
        ws.column_dimensions["A"].width = 45
        for ci in range(len(years)):
            ws.column_dimensions[chr(66+ci)].width = 18
        ws.freeze_panes = "B3"
    path = f"/tmp/{company_id}.xlsx"
    wb.save(path)
    return send_file(path, as_attachment=True,
        download_name=f"{c['code']}_KAP_{datetime.now().strftime('%Y%m%d')}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.route("/download/<company_id>/word")
def download_word(company_id):
    company = db_query("SELECT * FROM companies WHERE id=?", (company_id,))
    if not company:
        abort(404)
    c = company[0]
    rows = db_query("""
        SELECT year, table_type, label, value FROM financials
        WHERE company_id=? ORDER BY table_type, label, year DESC
    """, (company_id,))
    if not rows:
        abort(404)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = section.right_margin = Cm(1.5)
    h = doc.add_heading(f"{c['code']} | {c['title']}", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"KAP Finansal Veriler - {datetime.now().strftime('%d.%m.%Y')}")
    tables = {}
    years_set = set()
    for row in rows:
        years_set.add(row["year"])
        tables.setdefault(row["table_type"], {}).setdefault(
            row["label"].strip("/"), {})[row["year"]] = row["value"]
    years = sorted(years_set, reverse=True)
    for table_type, data in tables.items():
        doc.add_heading(table_type.upper(), level=1)
        table = doc.add_table(rows=1, cols=1+len(years))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Kalem"
        for ci, year in enumerate(years, 1):
            hdr[ci].text = str(year)
        for label, year_vals in sorted(data.items()):
            row_cells = table.add_row().cells
            row_cells[0].text = label
            for ci, year in enumerate(years, 1):
                val = year_vals.get(year)
                if val is not None:
                    row_cells[ci].text = f"{val:,.2f}"
        doc.add_paragraph()
    path = f"/tmp/{company_id}.docx"
    doc.save(path)
    return send_file(path, as_attachment=True,
        download_name=f"{c['code']}_KAP_{datetime.now().strftime('%Y%m%d')}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/api/stats")
def api_stats():
    stats = {
        "companies": db_query("SELECT COUNT(*) as n FROM companies")[0]["n"],
        "financials": db_query("SELECT COUNT(*) as n FROM financials")[0]["n"],
        "last_update": db_query("SELECT MAX(updated_at) as t FROM companies")[0]["t"],
        "years": [r["year"] for r in db_query("SELECT DISTINCT year FROM financials ORDER BY year DESC")],
    }
    return jsonify(stats)

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
